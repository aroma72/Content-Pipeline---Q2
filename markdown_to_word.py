"""
Convert Systems Evaluations Instructor Pack markdown to Word documents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    """Add a paragraph with optional formatting"""
    p = doc.add_paragraph(text)
    if bold or italic or size != 11:
        for run in p.runs:
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if size != 11:
                run.font.size = Pt(size)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font"""
    p = doc.add_paragraph(code_text, style='List Number')
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)

    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_table(doc, rows, cols, data):
    """Add a table to the document"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    # Fill in the data
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)

def convert_instructor_pack_to_word():
    """Convert Systems Evaluations instructor pack to Word"""

    # Read markdown file
    with open(r'c:\Users\Aroma Tahir\Downloads\Content Queen\weekly_artifacts\week-20-2026\act_outputs\systems_evaluations_instructor_pack.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Instructor Pack — Systems Evaluations', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    meta = doc.add_paragraph('Course: Agentic AI | Week: 20 | Date: 2026-05-11\nUnit ID: unit_systems_evaluations_w20 | Time Box: 90 minutes')
    meta.paragraph_format.space_before = Pt(6)
    meta.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # Spacer

    # Parse markdown and add to document
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('# '):
            if line.startswith('## '):
                add_heading(doc, line.replace('## ', ''), 2)
            elif line.startswith('# '):
                add_heading(doc, line.replace('# ', ''), 1)
            i += 1

        # Code blocks
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1

            code_block = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            p_format.right_indent = Inches(0.5)
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)

            run = p.add_run(code_block)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

            # Add background color to code block
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F2F2F2')
            p._element.get_or_add_pPr().append(shading_elm)

            i += 1

        # Tables (markdown format: |---|---|)
        elif line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                table_rows.append(row)
                i += 1

            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data.replace('---', '').strip()

        # Bullet points
        elif line.strip().startswith('-') or line.strip().startswith('*'):
            bullet_text = line.strip().lstrip('-*').strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1

        # Bold text (** ** or __ __)
        elif '**' in line or '__' in line:
            p = doc.add_paragraph()

            # Replace markdown formatting
            parts = re.split(r'(\*\*|__)', line)
            bold = False

            for part in parts:
                if part in ('**', '__'):
                    bold = not bold
                elif part:
                    run = p.add_run(part)
                    if bold:
                        run.font.bold = True

            i += 1

        # Regular paragraphs
        elif line.strip() and not line.startswith('#') and not line.startswith('---'):
            if line.strip():
                doc.add_paragraph(line.strip())
            i += 1

        # Empty lines
        else:
            i += 1

    # Save document
    output_path = r'c:\Users\Aroma Tahir\Downloads\Content Queen\weekly_artifacts\week-20-2026\Systems_Evaluations_Instructor_Pack.docx'
    doc.save(output_path)
    return output_path

def convert_learner_pack_to_word():
    """Convert Systems Evaluations learner pack to Word"""

    # Read markdown file
    with open(r'c:\Users\Aroma Tahir\Downloads\Content Queen\weekly_artifacts\week-20-2026\act_outputs\systems_evaluations_learner_pack.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Session Summary — Systems Evaluations', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    meta = doc.add_paragraph('Course: Agentic AI | Week: 20 | Date: 2026-05-11\nDuration: 90 minutes | Format: Lecture + Live Build + Hands-on Lab')
    meta.paragraph_format.space_before = Pt(6)
    meta.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # Spacer

    # Parse markdown and add to document
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('# '):
            if line.startswith('## '):
                add_heading(doc, line.replace('## ', ''), 2)
            elif line.startswith('# '):
                add_heading(doc, line.replace('# ', ''), 1)
            i += 1

        # Code blocks
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1

            code_block = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            p_format.right_indent = Inches(0.5)
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)

            run = p.add_run(code_block)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

            # Add background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F2F2F2')
            p._element.get_or_add_pPr().append(shading_elm)

            i += 1

        # Tables
        elif line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                table_rows.append(row)
                i += 1

            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data.replace('---', '').strip()

        # Bullet points
        elif line.strip().startswith('-') or line.strip().startswith('*'):
            bullet_text = line.strip().lstrip('-*').strip()
            if bullet_text:
                p = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1

        # Numbered lists
        elif line.strip() and line.strip()[0].isdigit() and '.' in line:
            num_text = line.strip()
            p = doc.add_paragraph(num_text.split('.', 1)[1].strip(), style='List Number')
            i += 1

        # Regular paragraphs
        elif line.strip() and not line.startswith('#') and not line.startswith('---'):
            if line.strip():
                doc.add_paragraph(line.strip())
            i += 1

        # Empty lines
        else:
            i += 1

    # Save document
    output_path = r'c:\Users\Aroma Tahir\Downloads\Content Queen\weekly_artifacts\week-20-2026\Systems_Evaluations_Learner_Pack.docx'
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    print("Converting Systems Evaluations materials to Word format...")

    instructor_path = convert_instructor_pack_to_word()
    print("Instructor Pack created: " + instructor_path)

    learner_path = convert_learner_pack_to_word()
    print("Learner Pack created: " + learner_path)

    print("\nBoth documents ready for distribution!")
