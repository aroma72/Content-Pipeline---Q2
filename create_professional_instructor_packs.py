"""
Create truly professional Word documents for instructor packs
Proper formatting - no markdown syntax visible
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
PRIMARY = RGBColor(31, 78, 121)      # Dark blue
ACCENT = RGBColor(0, 176, 176)       # Teal
LIGHT_GRAY = RGBColor(245, 248, 250)
TEXT_DARK = RGBColor(33, 33, 33)
TEXT_MED = RGBColor(97, 97, 97)

def shade_cell(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_professional_doc(title, subtitle=""):
    """Create professional document template"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY

    # Subtitle
    if subtitle:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = ACCENT
        subtitle_run.font.bold = True

    # Divider
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(12)

    return doc

def add_section_heading(doc, text):
    """Add section heading"""
    heading = doc.add_paragraph()
    heading.style = 'Heading 1'
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(8)

    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.text = text
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = PRIMARY

    return heading

def add_subsection_heading(doc, text):
    """Add subsection heading"""
    heading = doc.add_paragraph()
    heading.style = 'Heading 2'
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)

    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.text = text
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    return heading

def add_body_text(doc, text, size=11, bold=False, space_after=6):
    """Add formatted body text"""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(space_after)

    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = TEXT_DARK

    return para

def create_enhanced_instructor_pack():
    """Create enhanced instructor pack as professional Word document"""
    doc = create_professional_doc(
        "Systems Evaluations",
        "Instructor Pack - Enhanced Edition"
    )

    # Metadata
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run("Agentic AI Course | Week 20 | 90 Minutes\nUnit: unit_systems_evaluations_w20\nFramework: Scaffolding + TPACK + Project-Based Learning")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = TEXT_MED
    meta_run.italic = True
    meta_para.paragraph_format.space_after = Pt(14)

    # ===== LEARNING OBJECTIVES =====
    add_section_heading(doc, "Learning Objectives")

    add_body_text(doc, "By the end of this 90-minute session, learners will:")

    objectives = [
        "Define systems evaluation (vs testing) and articulate why it matters for production AI",
        "Design a 3-5 metric evaluation framework for a specific agent use case",
        "Build a hybrid evaluation suite combining automated checks and human review",
        "Apply scaffolded thinking: start with human judgment, then derive metrics",
        "Reflect on ethical dimensions: bias, fairness, and documentation in evaluation"
    ]

    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Number')
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # ===== TEACHING BRIEF =====
    add_section_heading(doc, "Teaching Brief")

    add_subsection_heading(doc, "What Learners Already Know")
    knowledge_items = [
        "How agents work (perceive → act → observe loop)",
        "How to build MCP servers and connect Claude",
        "Basic metrics (accuracy, precision, recall)",
        "Python and testing frameworks",
        "Drawing Room pipeline (signal → content → publish)"
    ]
    for item in knowledge_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ===== WEAK SPOTS TABLE =====
    add_subsection_heading(doc, "Critical Weak Spots to Monitor")

    weak_table = doc.add_table(rows=6, cols=4)
    weak_table.style = 'Light Grid Accent 1'

    # Header
    header_cells = weak_table.rows[0].cells
    headers = ["Issue", "Why It Happens", "What To Do", "Red Flag Signal"]
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        shade_cell(header_cells[i], "1F4E79")
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data
    weak_spots_data = [
        ("Confusing evaluation with testing",
         "They think testing = unit tests",
         "Use Quality Control analogy. Testing checks code; evaluation checks value",
         "They ask 'Is this a test?' when you say evaluation"),

        ("Metric overload",
         "Fear of missing visibility",
         "Say: Pick 3-5 that matter for THIS system. Ignore the rest.",
         "They want to track everything"),

        ("Not knowing where to start",
         "No mental model for design",
         "Start with human eval first. Then derive metrics from that.",
         "They ask 'What metrics?' before defining what good looks like"),

        ("Treating evaluation as one-time",
         "Waterfall thinking",
         "Emphasize continuous: measure → improve → measure again",
         "They expect a single evaluation moment"),

        ("Assuming all agents are the same",
         "Generic thinking",
         "Explicit: Design for YOUR system, not off-the-shelf",
         "They apply checklist metrics to everything")
    ]

    for i, (issue, why, what, signal) in enumerate(weak_spots_data, 1):
        cells = weak_table.rows[i].cells
        cells[0].text = issue
        cells[1].text = why
        cells[2].text = what
        cells[3].text = signal

        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ===== SESSION PLAN =====
    add_section_heading(doc, "Session Plan (90 Minutes)")

    session_table = doc.add_table(rows=7, cols=4)
    session_table.style = 'Light Grid Accent 1'

    session_headers = session_table.rows[0].cells
    session_headers[0].text = "Time"
    session_headers[1].text = "Block"
    session_headers[2].text = "Activity"
    session_headers[3].text = "Purpose"

    for cell in session_headers:
        shade_cell(cell, "00B0B0")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    session_data = [
        ("0:00–0:10", "Warm-up", "Show agent output. Ask: Is this good?", "Surface misconceptions"),
        ("0:10–0:25", "Concept 1", "What is Evaluation? MEASURE Loop", "Build mental model"),
        ("0:25–0:40", "Concept 2", "Types of Evaluation (Automated, Manual, Hybrid)", "Establish tradeoffs"),
        ("0:40–0:65", "Live Build", "Code 3 evaluators live (narrate thinking)", "Model procedural knowledge"),
        ("0:65–0:80", "Exercise", "Design rubric + checks + metrics (pairs)", "Guided practice with support"),
        ("0:80–0:90", "Reflect", "How do you know agent is good? Discussion", "Metacognition + iteration mindset")
    ]

    for i, (time, block, activity, purpose) in enumerate(session_data, 1):
        cells = session_table.rows[i].cells
        cells[0].text = time
        cells[1].text = block
        cells[2].text = activity
        cells[3].text = purpose

        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ===== EXPLANATION VARIANTS =====
    add_section_heading(doc, "Explanation Variants (Multiple Entry Points)")

    add_subsection_heading(doc, "Concept 1: What is Systems Evaluation?")

    variants = [
        ("VARIANT A: Quality Control Analogy (START HERE)",
         "Systems evaluation is to agents what quality control is to factories. You build a product, then ask: Does it work? Does it meet standards? What breaks? Why it works: Concrete, relatable, grounds abstract concept in physical world."),

        ("VARIANT B: Problem-First",
         "You've built an agent. But how do you know it's good? Evaluation answers: Does it generate quality materials? Does it publish without errors? Does it handle edge cases? Why it works: Addresses pain point. Shows relevance. Practical, not theoretical."),

        ("VARIANT C: Code-First",
         "Evaluation is a layer on your agent. Capture inputs → run → measure quality → iterate. Measure → find problems → improve → measure. Why it works: Maps to engineering mindset. Emphasizes feedback loops."),

        ("VARIANT D: Ethical Angle",
         "When you ship an agent, you're responsible. Evaluation means: Is it accurate? Fair? Can you explain decisions? Why it works: Connects to values. Shows evaluation as responsibility, not checkbox.")
    ]

    for variant_name, variant_text in variants:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

        title_run = p.add_run(variant_name)
        title_run.bold = True
        title_run.font.color.rgb = PRIMARY

        p.add_run("\n" + variant_text)

    doc.add_paragraph()

    # ===== FACILITATION NOTES =====
    add_section_heading(doc, "Facilitation Notes")

    add_subsection_heading(doc, "Warm-up (0:00-0:10)")

    facilitation_steps = [
        "Show a real agent response on screen",
        "Ask: 'Is this good? Rate it 1-5' (force criteria definition)",
        "Listen: Don't correct. Let them name what makes it good",
        "Synthesize: 'You said accurate, concise, complete. Those are metrics.'",
        "Insight: 'We just did evaluation. Now let's formalize it.'"
    ]

    for step in facilitation_steps:
        doc.add_paragraph(step, style='List Number')

    add_body_text(doc, "If they struggle: 'Imagine you're a learner. Would you understand this output?'")

    doc.add_paragraph()

    # ===== EXERCISE =====
    add_section_heading(doc, "Hands-On Exercise (0:65-0:80)")

    add_body_text(doc, "Scenario: Design evaluation for a Code Reviewer Agent\n", bold=True)

    exercise_parts = [
        ("Part 1: Human Evaluation (10 min)",
         "Read the rubric. Evaluate 3 sample code reviews. Rate on: Relevance, Clarity, Actionability (1-4 scale). Discuss: What makes a good review?"),

        ("Part 2: Design Automated Checks (15 min)",
         "Write 2 checks that verify basic quality. Example: Does review suggest at least 1 improvement? You design 2 more."),

        ("Part 3: Design Metrics (15 min)",
         "Pick 3 metrics to track daily. For each: What's the target? Example: % of reviews rated Relevant by humans (target: 85%+)"),

        ("Part 4: Share & Discuss (5 min)",
         "Each pair shares 1 check + 1 metric. Group votes: Which are most important?")
    ]

    for part_title, part_content in exercise_parts:
        p = doc.add_paragraph()
        p_run = p.add_run(part_title)
        p_run.bold = True
        p_run.font.color.rgb = ACCENT
        p.add_run("\n" + part_content)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ===== ASSESSMENT =====
    add_section_heading(doc, "Assessment Checkpoints")

    checkpoints = [
        ("Checkpoint 1 (0:25): Conceptual Understanding",
         "Ask: What's the difference between a unit test and evaluation?\nAcceptable: Unit test checks code works. Evaluation checks if system achieves its goal."),

        ("Checkpoint 2 (0:40): Framework Understanding",
         "Show code. Ask: What's this checking? Why?\nGreen flag: They explain both the logic AND the reasoning."),

        ("Checkpoint 3 (0:75): Application",
         "Watch their exercise. Can they write a sensible check or metric?\nGreen: They ask 'What would a human do first?'"),

        ("Checkpoint 4 (0:90): Reflection",
         "Ask: Is evaluation one-time or ongoing?\nAcceptable: Ongoing. You measure, improve, measure again.")
    ]

    for cp_title, cp_content in checkpoints:
        p = doc.add_paragraph()
        p_run = p.add_run(cp_title)
        p_run.bold = True
        p_run.font.color.rgb = PRIMARY
        p.add_run("\n" + cp_content)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # ===== ETHICS =====
    add_section_heading(doc, "Ethical Considerations")

    ethics_items = [
        ("E1: Evaluation Encodes Values",
         "Every metric reflects a value judgment. 'Accuracy' assumes correctness matters most. But speed or empathy might matter more.\n\nIn Practice: Name your values explicitly. 'For this system, we prioritize accuracy over speed because...'"),

        ("E2: Test Sets Are Biased",
         "Evaluation is only as good as your test data. If test set doesn't include edge cases, minorities, or unusual scenarios, it will miss failures.\n\nIn Practice: Stratify test set (80% typical cases, 10% edge cases, 10% underrepresented groups)."),

        ("E3: Metric Gaming",
         "If you optimize for a metric, humans will game it. 'Maximize response length' → agent writes novels instead of concise answers.\n\nIn Practice: Monitor not just the metric you optimize, but proxy metrics (user satisfaction, latency, cost)."),

        ("E4: Documentation & Audit Trail",
         "You decide to pass an agent. Why? If you can't explain, you're not accountable.\n\nIn Practice: Log what was measured, why, what the threshold was, who decided, when.")
    ]

    for ethics_title, ethics_content in ethics_items:
        p = doc.add_paragraph()
        p_run = p.add_run(ethics_title)
        p_run.bold = True
        p_run.font.color.rgb = ACCENT
        p.add_run("\n" + ethics_content)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # ===== RESOURCES =====
    add_section_heading(doc, "Resources & Materials")

    resources = [
        ("Slides",
         "Systems_Evaluations_Instructor_Slides_Final.pptx (12 slides, all concepts)\nSystems_Evaluations_Instructor_Slides_Final.pdf (shareable version)"),

        ("Code Examples",
         "All Python examples in this pack are executable. Available in GitHub repository."),

        ("Handouts to Print",
         "1. Rubric Guide (1 page)\n2. Metric Worksheet (1 page)\n3. Check Checklist (1 page)\n4. Ethical Audit (1 page)"),

        ("Further Reading",
         "• ACUE: 10 Best Practices for AI Assignments in Higher Ed\n• The TPACK Framework Explained\n• 7 Scaffolding Learning Strategies for the Classroom")
    ]

    for resource_title, resource_content in resources:
        p = doc.add_paragraph()
        p_run = p.add_run(resource_title)
        p_run.bold = True
        p_run.font.size = Pt(12)
        p_run.font.color.rgb = PRIMARY
        p.add_run("\n" + resource_content)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # ===== FOOTER =====
    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(12)
    footer.paragraph_format.border_top = True

    footer_run = footer.add_run("Last Updated: 2026-05-07 | Version: Enhanced (v2.0)\naroma.tahir@taleemabad.com")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = TEXT_MED
    footer_run.italic = True

    return doc

def create_original_instructor_pack():
    """Create original instructor pack"""
    doc = create_professional_doc(
        "Systems Evaluations",
        "Instructor Pack - Core Edition"
    )

    # Metadata
    meta = doc.add_paragraph()
    meta_run = meta.add_run("Agentic AI | Week 20 | 90 Minutes")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = TEXT_MED
    meta.paragraph_format.space_after = Pt(12)

    # Teaching Brief
    add_section_heading(doc, "Teaching Brief")

    add_subsection_heading(doc, "What Learners Already Know")
    knowledge = [
        "How agents work (perceive → act → observe loop)",
        "How to build MCP servers and connect Claude",
        "Basic metrics (accuracy, precision, recall)",
        "Python and testing frameworks",
        "Drawing Room pipeline"
    ]
    for item in knowledge:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Weak Spots
    add_subsection_heading(doc, "Watch Out For (Common Misconceptions)")

    misconceptions = [
        "Confusing evaluation with testing (evaluation is broader than unit tests)",
        "Metric overload (they'll want to measure everything)",
        "Not knowing where to start (start with human eval, then derive metrics)",
        "Treating evaluation as one-time (it's continuous)",
        "Assuming all agents are the same (metrics are system-specific)"
    ]

    for item in misconceptions:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Session Plan
    add_section_heading(doc, "Session Plan (90 Minutes)")

    plan_table = doc.add_table(rows=7, cols=3)
    plan_table.style = 'Light Grid Accent 1'

    plan_headers = plan_table.rows[0].cells
    plan_headers[0].text = "Time"
    plan_headers[1].text = "Block"
    plan_headers[2].text = "Activity"

    for cell in plan_headers:
        shade_cell(cell, "1F4E79")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    plan_data = [
        ("0:00–0:10", "Warm-up", "Show agent output. Is this good?"),
        ("0:10–0:25", "Concept", "What is Evaluation? MEASURE Loop"),
        ("0:25–0:45", "Concept", "Types of Evaluation & Metrics"),
        ("0:45–0:65", "Live Build", "Build evaluation suite (live code)"),
        ("0:65–0:80", "Task", "Design metrics for scenario"),
        ("0:80–0:90", "Reflect", "How to know when agent is good enough?")
    ]

    for i, (time, block, activity) in enumerate(plan_data, 1):
        cells = plan_table.rows[i].cells
        cells[0].text = time
        cells[1].text = block
        cells[2].text = activity

    doc.add_paragraph()

    # Explanation Variants
    add_section_heading(doc, "Explanation Variants")

    add_subsection_heading(doc, "Concept: What is Systems Evaluation?")

    variants = [
        ("Quality Control Analogy (START HERE)",
         "Systems evaluation is to agents what quality control is to factories. You build a product, ask: Does it work? Does it meet standards? What breaks? Evaluation answers these."),

        ("Problem-First",
         "You built an agent. How do you know it's good? Evaluation answers: Does it generate quality materials? Does it publish correctly? Does it handle edge cases?"),

        ("Code-First",
         "Evaluation is a layer on your agent. Capture inputs → run agent → measure quality → iterate. It's continuous validation, not one-time testing.")
    ]

    for variant_name, variant_text in variants:
        p = doc.add_paragraph()
        p_run = p.add_run(variant_name)
        p_run.bold = True
        p_run.font.color.rgb = ACCENT
        p.add_run("\n" + variant_text)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # Quick Reference
    add_section_heading(doc, "Key Concepts")

    concepts = {
        "MEASURE Loop": "Measure → Explore → Assess → Sum → Unblock (continuous improvement cycle)",
        "Automation": "Fast, scalable, cheap, limited nuance",
        "Human Eval": "Nuanced, context-aware, slow, expensive",
        "Hybrid": "Automated first pass + human spot-check (best approach)"
    }

    for concept, definition in concepts.items():
        p = doc.add_paragraph()
        p_run = p.add_run(concept + ": ")
        p_run.bold = True
        p.add_run(definition)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer_run = footer.add_run("aroma.tahir@taleemabad.com")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = TEXT_MED

    return doc

# Create documents
print("\nCreating professional instructor pack documents...\n")

print("[1/2] Enhanced Pack...")
enhanced = create_enhanced_instructor_pack()
enhanced.save("weekly_artifacts/week-20-2026/act_outputs/Systems_Evaluations_Instructor_Pack_ENHANCED.docx")
print("      Done. Ready to share.")

print("[2/2] Original Pack...")
original = create_original_instructor_pack()
original.save("weekly_artifacts/week-20-2026/act_outputs/Systems_Evaluations_Instructor_Pack_ORIGINAL.docx")
print("      Done. Quick reference version.")

print("\nAll documents created successfully!")
print("\nBoth files are now:")
print("  ✓ Professional formatting (no markdown visible)")
print("  ✓ Proper tables, bullets, and styling")
print("  ✓ Ready to print or share")
print("  ✓ Professionally formatted typography\n")
