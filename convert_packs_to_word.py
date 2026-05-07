"""
Convert Instructor Pack markdown files to professional Word documents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styled_document(title, subtitle=""):
    """Create a new Document with professional styling"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)  # Dark blue

    if subtitle:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(0, 176, 176)  # Teal

    # Add horizontal line
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    return doc

def add_heading(doc, text, level=1):
    """Add formatted heading"""
    if level == 1:
        para = doc.add_paragraph(text, style='Heading 1')
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:
        para = doc.add_paragraph(text, style='Heading 2')
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
    else:
        para = doc.add_paragraph(text, style='Heading 3')
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)

def add_section_box(doc, title, content):
    """Add a highlighted section box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)

    title_run = para.add_run(title)
    title_run.font.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    para.add_run("\n" + content)
    para.paragraph_format.space_after = Pt(8)

def convert_enhanced_pack_to_word():
    """Convert the enhanced instructor pack to Word"""
    doc = create_styled_document(
        "Systems Evaluations",
        "Instructor Pack - Enhanced Edition"
    )

    # Metadata
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run("Course: Agentic AI | Week 20 | 90 minutes\nUnit ID: unit_systems_evaluations_w20\nPedagogical Framework: Scaffolding + TPACK + Project-Based Learning")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    meta_para.paragraph_format.space_after = Pt(12)

    # Learning Objectives
    add_heading(doc, "Learning Objectives (SMART)", 1)
    doc.add_paragraph(
        "By the end of this 90-minute session, learners will:",
        style='List Number'
    )

    objectives = [
        "Define what systems evaluation means (vs testing) and articulate why it matters",
        "Design a 3-5 metric evaluation framework for a specific agent use case",
        "Build a hybrid evaluation suite combining automated checks + human review",
        "Apply scaffolded thinking: start with human judgment, then derive metrics",
        "Reflect on ethical dimensions: bias, fairness, and documentation"
    ]

    for obj in objectives:
        doc.add_paragraph(obj, style='List Number')

    doc.add_paragraph()

    # Teaching Brief
    add_heading(doc, "Teaching Brief", 1)

    add_heading(doc, "What Learners Already Know", 2)
    knowledge = [
        "How agents work (perceive → act → observe loop)",
        "How to build MCP servers and connect Claude",
        "Basic metrics (accuracy, precision, recall)",
        "Python and testing frameworks basics",
        "Drawing Room pipeline (signal → content → publish)"
    ]
    for item in knowledge:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    add_heading(doc, "Key Weak Spots to Monitor", 2)

    # Create weak spots table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    headers = ["Issue", "Why It Happens", "What To Do", "Signal to Watch"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    weak_spots = [
        ("Confusing evaluation with testing",
         "Mental model: testing = unit tests",
         "Use Quality Control analogy. Show: testing checks code; evaluation checks value",
         "They ask 'Is this a test?' when you say evaluation"),

        ("Metric overload",
         "Fear of missing something",
         "Say 'Pick 3-5 that matter for THIS system'",
         "They want to track everything"),

        ("Not knowing where to start",
         "No mental model for evaluation design",
         "Start with human eval first, then derive metrics",
         "They ask 'What metrics?' before knowing what good looks like"),

        ("Treating evaluation as one-time",
         "Waterfall thinking",
         "Emphasize continuous loop: measure → improve → measure",
         "They ask 'When do we evaluate?' expecting single answer"),

        ("Assuming all agents are the same",
         "Generic metric thinking",
         "Explicit: Design for YOUR system, not generic",
         "They try to apply checklist metrics to new domains")
    ]

    for i, (issue, why, what, signal) in enumerate(weak_spots, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = issue
        row_cells[1].text = why
        row_cells[2].text = what
        row_cells[3].text = signal

    doc.add_paragraph()

    # Session Plan
    add_heading(doc, "Session Plan (90 Minutes)", 1)

    session_plan = [
        ("0:00–0:10", "Warm-up", "Show agent output. Ask: Is this good? How do you know?"),
        ("0:10–0:25", "Concept: Evaluation Framework", "MEASURE loop + why it matters"),
        ("0:25–0:40", "Concept: Types of Evaluation", "Automated checks, metrics, human eval"),
        ("0:40–0:65", "Live Build", "Code 3 evaluators (model + guidance)"),
        ("0:65–0:80", "Exercise", "Design rubric + checks + metrics for scenario"),
        ("0:80–0:90", "Reflect & Close", "How do you know when your agent is good?")
    ]

    plan_table = doc.add_table(rows=len(session_plan)+1, cols=3)
    plan_table.style = 'Light Grid Accent 1'

    plan_header = plan_table.rows[0].cells
    plan_header[0].text = "Time"
    plan_header[1].text = "Block"
    plan_header[2].text = "Activity"
    for cell in plan_header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for i, (time, block, activity) in enumerate(session_plan, 1):
        row = plan_table.rows[i].cells
        row[0].text = time
        row[1].text = block
        row[2].text = activity

    doc.add_paragraph()

    # Explanation Variants
    add_heading(doc, "Explanation Variants (Multiple Entry Points)", 1)

    add_heading(doc, "Concept 1: What is Systems Evaluation?", 2)

    add_section_box(doc, "Variant A — Quality Control Analogy (START HERE)",
        "Systems evaluation is to agents what quality control is to factories. You build a product (an agent), then you ask: Does it work? Does it meet standards? What breaks? Evaluation answers these questions.\n\nWhy it works: Concrete, relatable, anchors abstract concept.")

    add_section_box(doc, "Variant B — Problem-first",
        "You've built an agent. But how do you know it's good? Good is subjective. Evaluation answers: Does it generate quality materials? Does it publish without errors? Does it handle edge cases? It's the difference between 'works on my test' and 'works in production'.")

    add_section_box(doc, "Variant C — Code-first",
        "Evaluation is a layer on your agent. Capture inputs → run agent → measure quality → iterate. Measure → find problems → improve → measure again. It's continuous validation.")

    add_section_box(doc, "Variant D — Ethical angle",
        "When you ship an agent, you're responsible. Evaluation means: Is it accurate? Is it fair? Can you explain decisions? Does it handle edge cases? These aren't technical—they're ethical.")

    doc.add_paragraph()

    # Facilitation Notes
    add_heading(doc, "Facilitation Notes (How to Lead This)", 1)

    add_heading(doc, "Warm-up (0:00-0:10): Make It Interactive", 2)
    facilitation_text = """
1. Show a real agent response on screen
2. Ask: "Is this good? Rate it 1-5" (force them to define criteria)
3. Listen: Don't correct. Let them name criteria
4. Synthesize: "You said accurate, concise, complete. Those are your metrics"
5. Insight: "We just did evaluation without knowing it. Now let's formalize"

If they struggle: "Imagine you're a learner. Would you understand this?"
"""
    doc.add_paragraph(facilitation_text)

    doc.add_paragraph()

    # Hands-on Exercise
    add_heading(doc, "Hands-On Exercise (0:65-0:80)", 1)

    exercise_text = """
Scenario: Design evaluation for a Code Reviewer Agent

Part 1: Human Evaluation (10 min)
- Read the rubric
- Evaluate 3 sample code reviews
- Rate on: Relevance, Clarity, Actionability (1-4)
- Discuss: What makes a good review?

Part 2: Design Automated Checks (15 min)
- Write 2 checks that verify basic quality
- Example: "Does review suggest at least 1 improvement?"
- You design 2 more

Part 3: Design Metrics (15 min)
- Pick 3 metrics to track daily
- For each: What's the target?
- Example: "% of reviews rated Relevant by humans (target: 85%+)"

Part 4: Share & Discuss (5 min)
- Each pair shares 1 check + 1 metric
- Group votes: Most important?
"""
    doc.add_paragraph(exercise_text)

    doc.add_paragraph()

    # Assessment Checkpoints
    add_heading(doc, "Assessment Checkpoints", 1)

    checkpoints = [
        ("Checkpoint 1 (0:25)", "Conceptual Understanding",
         "Ask: What's the difference between a unit test and an evaluation?\nAcceptable: Unit test checks code; evaluation checks if system achieves its goal"),

        ("Checkpoint 2 (0:40)", "Framework Understanding",
         "Show code. Ask: What's this checking? Why?\nGreen flag: They can point to logic AND explain reasoning"),

        ("Checkpoint 3 (0:75)", "Application",
         "Observe exercise: Can they write a sensible check or metric?\nGreen: They ask 'What would a human do first?' (shows understanding)"),

        ("Checkpoint 4 (0:90)", "Reflection",
         "Ask: Is evaluation one-time or ongoing?\nAcceptable: Ongoing. Measure, improve, measure again")
    ]

    for title, checkpoint_type, content in checkpoints:
        add_heading(doc, f"{title}: {checkpoint_type}", 2)
        doc.add_paragraph(content)

    doc.add_paragraph()

    # Ethical Considerations
    add_heading(doc, "Ethical Considerations (Critical for AI)", 1)

    ethics = [
        ("E1: Evaluation Encodes Values",
         "Every metric reflects values. 'Accuracy' assumes correctness matters most. But speed or empathy might matter more.\nIn Practice: Name your values explicitly. 'For this system, we prioritize accuracy over speed because...'"),

        ("E2: Test Sets Are Biased",
         "Evaluation is only as good as test data. If test set doesn't include edge cases, minorities, or unusual scenarios, it will miss failures.\nIn Practice: Stratify test set (80% typical, 10% edge cases, 10% underrepresented groups)"),

        ("E3: Metric Gaming",
         "If you optimize for a metric, humans will game it. Maximize response length → agent writes novels.\nIn Practice: Monitor not just the metric, but proxy metrics (satisfaction, latency, cost)"),

        ("E4: Documentation & Audit Trail",
         "You decide to pass an agent. Why? If you can't explain, you're not accountable.\nIn Practice: Log what was measured, why, threshold, who decided, when")
    ]

    for title, content in ethics:
        add_heading(doc, title, 2)
        doc.add_paragraph(content)

    # Resources
    add_heading(doc, "Resources & Materials", 1)

    resources_text = """
Slides:
• Systems_Evaluations_Instructor_Slides_Final.pptx (12 slides, all concepts)
• Systems_Evaluations_Instructor_Slides_Final.pdf (shareable)

Code Examples:
• All examples in this pack are executable Python
• Available in GitHub repository for copy-paste

Further Reading:
• ACUE: 10 Best Practices for AI Assignments in Higher Ed
• The TPACK Framework Explained
• 7 Scaffolding Learning Strategies for the Classroom
• Pedagogical Content Knowledge

Handouts to Print:
1. Rubric Guide (1 page)
2. Metric Worksheet (1 page)
3. Check Checklist (1 page)
4. Ethical Audit (1 page)
"""

    doc.add_paragraph(resources_text)

    doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.space_before = Pt(12)
    footer_run = footer_para.add_run("Last Updated: 2026-05-07 | Version: Enhanced with Best Practices (v2.0)\nContact: aroma.tahir@taleemabad.com")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.italic = True

    return doc

def convert_original_pack_to_word():
    """Convert the original instructor pack to Word"""
    doc = create_styled_document(
        "Systems Evaluations",
        "Instructor Pack - Core Edition"
    )

    # Metadata
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run("Course: Agentic AI | Week 20 | 90 minutes\nUnit ID: unit_systems_evaluations_w20")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    meta_para.paragraph_format.space_after = Pt(12)

    # Teaching Brief
    add_heading(doc, "Teaching Brief", 1)

    add_heading(doc, "What Learners Already Know", 2)
    knowledge = [
        "How agents work (perceive → act → observe loop)",
        "How to build MCP servers and connect Claude",
        "Basic metrics (accuracy, precision, recall)",
        "What 'good output' means for their domain",
        "Python and testing frameworks basics"
    ]
    for item in knowledge:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    add_heading(doc, "Likely Weak Spots", 2)
    weak_spots_text = """
1. Confusing evaluation with testing — They'll think testing = unit tests. Evaluation is broader.
2. Metric overload — They want to measure everything. Be clear: pick 3-5 that matter.
3. Not knowing where to start — They'll ask "how do I measure Claude?" Answer: human eval first.
4. Treating evaluation as one-time — Evaluation is continuous. Measure → tweak → measure.
5. Assuming all agents are the same — Remind them: Drawing Room metrics ≠ chatbot metrics.
"""
    doc.add_paragraph(weak_spots_text)

    doc.add_paragraph()

    # Session Plan
    add_heading(doc, "Session Plan (90 Minutes)", 1)

    session_table = doc.add_table(rows=7, cols=3)
    session_table.style = 'Light Grid Accent 1'

    header_cells = session_table.rows[0].cells
    header_cells[0].text = "Time"
    header_cells[1].text = "Block"
    header_cells[2].text = "Activity"

    session_data = [
        ("0:00–0:10", "Warm-up", "Show agent output. Ask: Is this good?"),
        ("0:10–0:25", "Concept", "Evaluation Framework"),
        ("0:25–0:45", "Concept", "Metrics & Testing"),
        ("0:45–0:65", "Live Build", "Build evaluation suite"),
        ("0:65–0:80", "Task", "Design evaluation metrics"),
        ("0:80–0:90", "Reflect", "How to know when agent is good enough?")
    ]

    for i, (time, block, activity) in enumerate(session_data, 1):
        cells = session_table.rows[i].cells
        cells[0].text = time
        cells[1].text = block
        cells[2].text = activity

    doc.add_paragraph()

    # Explanation Variants
    add_heading(doc, "Explanation Variants", 1)

    add_heading(doc, "Concept 1: What is Systems Evaluation?", 2)

    variants = [
        ("Variant A — Quality Control Analogy (START HERE)",
         "Systems evaluation is to agents what quality control is to factories. You build a product, then ask: Does it work? Does it meet standards? What breaks? Evaluation answers those questions."),

        ("Variant B — Problem-first",
         "You've built an agent. But how do you know it's good? Good is subjective. Does it generate learning materials learners learn from? Does it publish without errors? Evaluation answers these."),

        ("Variant C — Code-first",
         "Evaluation is a layer on your agent. Capture inputs, run agent, measure quality against criteria. Do this repeatedly: measure → find problems → improve → measure again.")
    ]

    for variant_title, variant_content in variants:
        add_heading(doc, variant_title, 3)
        doc.add_paragraph(variant_content)

    doc.add_paragraph()

    # Code Examples
    add_heading(doc, "Example 1: Simple Pass/Fail Evaluation", 2)
    doc.add_paragraph("A minimal evaluator that checks if an agent output meets basic requirements.")

    code_example = """
def evaluate_learner_pack(pack: dict) -> EvaluationResult:
    \"\"\"Check if learner pack meets minimum standards.\"\"\"
    issues = []

    # Check: has all required sections
    required = ["session_summary", "glossary", "watch_order", "key_concepts"]
    for section in required:
        if section not in pack:
            issues.append(f"Missing section: {section}")

    # Check: glossary has minimum terms
    if "glossary" in pack:
        if len(pack["glossary"]) < 5:
            issues.append(f"Glossary too short")

    return EvaluationResult(
        passed=len(issues) == 0,
        issues=issues
    )
"""

    doc.add_paragraph(code_example)
    doc.add_paragraph()

    doc.add_paragraph("Teaching Point: Start simple. Don't build complexity you don't need.")

    doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.space_before = Pt(12)
    footer_run = footer_para.add_run("Last Updated: 2026-05-07\nContact: aroma.tahir@taleemabad.com")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.italic = True

    return doc

# Create both Word documents
print("\n" + "="*70)
print("CONVERTING INSTRUCTOR PACKS TO WORD FORMAT")
print("="*70 + "\n")

print("Creating Enhanced Pack...")
enhanced_doc = convert_enhanced_pack_to_word()
enhanced_path = "weekly_artifacts/week-20-2026/act_outputs/Systems_Evaluations_Instructor_Pack_ENHANCED.docx"
enhanced_doc.save(enhanced_path)
print(f"[OK] Enhanced Pack: {enhanced_path}")

print("Creating Original Pack...")
original_doc = convert_original_pack_to_word()
original_path = "weekly_artifacts/week-20-2026/act_outputs/Systems_Evaluations_Instructor_Pack_ORIGINAL.docx"
original_doc.save(original_path)
print(f"[OK] Original Pack: {original_path}")

print("\n" + "="*70)
print("COMPLETE! Both instructor packs converted to Word")
print("="*70)
print("\nFiles Created:")
print(f"  1. Systems_Evaluations_Instructor_Pack_ENHANCED.docx")
print(f"  2. Systems_Evaluations_Instructor_Pack_ORIGINAL.docx")
print("\nFeatures:")
print("  - Professional formatting (headers, styles, colors)")
print("  - Tables for session plan & weak spots")
print("  - Color-coded headings (navy blue)")
print("  - Proper spacing and typography")
print("  - Ready to print or share")
print("\n" + "="*70 + "\n")
